from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Title
title = doc.add_heading('Система автоматизации учета рабочего времени и планов', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Автоматизированный Telegram-бот с геймификацией и интеграцией с CrocoTime').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Section 1
doc.add_heading('1. Цель системы', level=1)
doc.add_paragraph('Снять с руководителя рутину по утреннему контролю команды, автоматизировать сбор ежедневных планов и учет прихода на работу, а также повысить дисциплину сотрудников с помощью встроенной системы геймификации.')

# Section 2
doc.add_heading('2. Используемые интеграции', level=1)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Telegram: ').bold = True
p.add_run('Удобный интерфейс для сотрудников (напоминания, отправка геопозиции и голосовых/текстовых планов). Групповой чат для оповещений.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Google Таблицы: ').bold = True
p.add_run('Центральная база данных (список сотрудников, графики начала работы, история чекинов и планы).')

p = doc.add_paragraph(style='List Bullet')
p.add_run('CrocoTime: ').bold = True
p.add_run('Автоматический фоновый контроль первой активности за рабочим ПК.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Google Web Speech API: ').bold = True
p.add_run('Автоматическая текстовая расшифровка голосовых сообщений.')

# Section 3
doc.add_heading('3. Логика работы (Сценарий рабочего дня)', level=1)

doc.add_heading('3.1. Индивидуальные напоминания (Пинг)', level=2)
doc.add_paragraph('У каждого сотрудника в Google Таблице задано свое индивидуальное время начала рабочего дня. Ровно за 10 минут до этого времени бот отправляет сотруднику персональное сообщение с запросом плана на день.')

doc.add_heading('3.2. Сбор ежедневных планов', level=2)
doc.add_paragraph('Сотрудник может ответить боту текстом или записать голосовое сообщение:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Бот самостоятельно расшифровывает аудио в текст.')
doc.add_paragraph('Голосовое сообщение и его текстовая расшифровка отправляются в личные сообщения руководителю/администратору.', style='List Bullet')
doc.add_paragraph('План автоматически заносится в Google Таблицу.', style='List Bullet')
doc.add_paragraph('(Планы в общую группу не пересылаются для сохранения конфиденциальности рабочих процессов).', style='List Bullet')

doc.add_heading('3.3. Фиксация начала рабочего дня', level=2)
doc.add_paragraph('Начало рабочего дня фиксируется двумя способами, в зависимости от формата работы сотрудника:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Офисные сотрудники (ПК): ').bold = True
p.add_run('Бот каждые 5 минут фоном обращается к серверу CrocoTime и автоматически фиксирует время первой активности за ПК. Сотруднику не нужно ничего нажимать.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Удаленные / Выездные сотрудники: ').bold = True
p.add_run('Бот запрашивает отправку геопозиции прямо в Telegram. Моментом начала рабочего дня считается время отправки локации.')
doc.add_paragraph('После успешной фиксации бот отправляет короткое оповещение в общую группу (например, "Александр начал работу").')

doc.add_heading('3.4. Контроль опозданий и Геймификация', level=2)
doc.add_paragraph('После чекина бот сверяет фактическое время старта с назначенным. Для мотивации сотрудников внедрена система "Стриков" (непрерывных дней без опозданий):')
doc.add_paragraph('Пришел вовремя: Счётчик дней подряд увеличивается на 1.', style='List Bullet')
doc.add_paragraph('Опоздал: Счётчик сбрасывается до 0.', style='List Bullet')

doc.add_paragraph('При достижении круглых дат дисциплины, бот публично награждает сотрудника в общей Telegram-группе:')
doc.add_paragraph('5 дней подряд — Красавчик', style='List Bullet')
doc.add_paragraph('20 дней подряд — Супер Красавчик', style='List Bullet')
doc.add_paragraph('60 дней подряд — Турбо Красавчик', style='List Bullet')
doc.add_paragraph('120 дней подряд — Супер Турбо Красавчик', style='List Bullet')
doc.add_paragraph('240 дней (год) подряд — Мега Красавчик', style='List Bullet')
doc.add_paragraph('Текущий прогресс и звания автоматически сохраняются в Google Таблицу.')

# Section 4
doc.add_heading('4. Главные преимущества для бизнеса', level=1)
doc.add_paragraph('1. Полное отсутствие микроменеджмента: Бот сам контролирует расписание каждого человека и напоминает о планах.')
doc.add_paragraph('2. Прозрачность данных: Вся история приходов, опозданий и задач формируется в аккуратной Google Таблице.')
doc.add_paragraph('3. Удобство для руководителя: Больше не нужно слушать длинные голосовые сообщения от подчиненных — бот пришлет готовый текст.')
doc.add_paragraph('4. Повышение дисциплины: Публичное признание успехов (система "Красавчиков") мотивирует команду приходить вовремя и не сбрасывать свои заработанные достижения.')

output_path = os.path.expanduser("~/Downloads/Workday_Bot_Documentation.docx")
doc.save(output_path)
print(f"Saved to {output_path}")

